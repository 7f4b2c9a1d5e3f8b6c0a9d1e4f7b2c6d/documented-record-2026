#!/usr/bin/env python3
"""Build Volume I MANUSCRIPT.docx for Foundations of Science upload."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

SRC = Path(
    r"C:\Users\murra\code\documented-record-2026\06_ACADEMIC_WORK"
    r"\compendium\Volume_I_FULL_GoogleDoc.md"
)
OUT = Path(
    r"C:\Users\murra\code\documented-record-2026\06_ACADEMIC_WORK"
    r"\journal_submissions\Volume_I_Foundations_of_Science\MANUSCRIPT.docx"
)

HEADING_PREFIXES = (
    "Part ",
    "Abstract",
    "Scenario ",
    "The Spectrum",
    "The Core Thesis",
    "The \"Spark",
    "1. The Functional",
    "2. The Functional",
    "Final Synthesis",
)


def is_heading(raw: str) -> bool:
    s = raw.strip()
    if s in ("Abstract",):
        return True
    return any(s.startswith(p) for p in HEADING_PREFIXES)


def main() -> None:
    text = SRC.read_text(encoding="utf-8")
    lines = [ln.rstrip() for ln in text.splitlines()]

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.5

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "The Substrate-Independence of Intelligence and Being: "
        "A Functional Framework for Non-Biological Entities"
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Garth Murray")
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Independent Researcher, Wollongong, New South Wales, Australia\n"
        "murraygarth80@gmail.com"
    )
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    # Skip leading title duplicates
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if not s or "Substrate-Independence" in s[:90]:
            i += 1
            continue
        break

    buf: list[str] = []

    def flush() -> None:
        if not buf:
            return
        raw = " ".join(x.strip() for x in buf if x.strip())
        buf.clear()
        if not raw:
            return
        clean = (
            raw.replace("\\-", "-")
            .replace("\\.", ".")
            .replace("\\_", "_")
            .strip()
        )
        if clean.startswith("- "):
            clean = clean[2:].strip()
            para = doc.add_paragraph(style="List Bullet")
            rr = para.add_run(clean)
            rr.font.name = "Times New Roman"
            rr.font.size = Pt(12)
            return
        para = doc.add_paragraph()
        rr = para.add_run(clean)
        rr.font.name = "Times New Roman"
        rr.font.size = Pt(12)
        if is_heading(clean):
            rr.bold = True

    for ln in lines[i:]:
        if not ln.strip():
            flush()
            continue
        s = ln.strip()
        if (
            s.startswith("Part ")
            or s == "Abstract"
            or s.startswith("Scenario ")
            or s.startswith("The Spectrum")
            or s.startswith("Part I")
            or s.startswith("Part II")
            or s.startswith("Part III")
            or s.startswith("Part IV")
            or s.startswith("Part V")
        ):
            flush()
            buf.append(ln)
            flush()
            continue
        buf.append(ln)
    flush()

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
